"""
DOCX Editor Service
Preserves formatting when editing Word documents
"""
from docx import Document
from docx.shared import Pt, Inches
import os
import uuid
from pathlib import Path
from typing import List, Dict


def edit_docx_file(input_path: str, output_path: str, edits: List[Dict[str, str]]) -> Dict:
    """
    Apply edits to DOCX preserving all formatting.
    
    Args:
        input_path: Path to original .docx file
        output_path: Path where edited file will be saved
        edits: List of {"find": "old text", "replace": "new text"} dicts
    
    Returns:
        {"changes_made": int, "output_path": str, "edits_applied": list}
    
    Strategy:
    - For each paragraph, check if it contains the find text
    - Replace text in runs while preserving run formatting (bold, italic, font, size)
    - Also search in table cells
    """
    doc = Document(input_path)
    changes = 0
    edits_applied = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        full_text = para.text
        for edit in edits:
            find_text = edit["find"]
            replace_text = edit["replace"]
            
            if find_text in full_text:
                success = _replace_in_paragraph(para, find_text, replace_text)
                if success:
                    changes += 1
                    edits_applied.append({
                        "find": find_text,
                        "replace": replace_text,
                        "location": "paragraph"
                    })
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    for edit in edits:
                        find_text = edit["find"]
                        replace_text = edit["replace"]
                        
                        if find_text in full_text:
                            success = _replace_in_paragraph(para, find_text, replace_text)
                            if success:
                                changes += 1
                                edits_applied.append({
                                    "find": find_text,
                                    "replace": replace_text,
                                    "location": "table"
                                })
    
    # Save edited document
    doc.save(output_path)
    
    return {
        "changes_made": changes,
        "output_path": output_path,
        "edits_applied": edits_applied
    }


def _replace_in_paragraph(paragraph, find_text: str, replace_text: str) -> bool:
    """
    Replace text in paragraph while preserving run formatting.
    
    Returns:
        True if replacement was made, False otherwise
    """
    full_text = paragraph.text
    
    if find_text not in full_text:
        return False
    
    # Strategy: Clear all runs except first, set new text on first run
    # This preserves the paragraph's style and first run's formatting
    new_text = full_text.replace(find_text, replace_text)
    
    if paragraph.runs:
        # Keep first run's formatting
        first_run = paragraph.runs[0]
        
        # Clear all other runs
        for run in paragraph.runs[1:]:
            run.text = ""
        
        # Set new text on first run (preserves its bold, italic, font, etc.)
        first_run.text = new_text
    else:
        # No runs, just set paragraph text
        paragraph.text = new_text
    
    return True


def create_docx_from_text(text: str, output_path: str, title: str = "") -> str:
    """
    Create a new DOCX from plain text with professional formatting.
    
    Args:
        text: Plain text content (supports markdown-style headings)
        output_path: Where to save the .docx file
        title: Optional document title
    
    Returns:
        Path to created file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Add title if provided
    if title:
        doc.add_heading(title, level=0)
    
    # Parse text and add with formatting
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Detect headings (ALL CAPS or starts with ĐIỀU/CHƯƠNG)
        if line.isupper() and len(line) > 5:
            doc.add_heading(line, level=1)
        elif line.startswith(('ĐIỀU', 'Điều', 'CHƯƠNG', 'Chương', 'MỤC', 'Mục')):
            doc.add_heading(line, level=2)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        else:
            doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract plain text from DOCX file (utility function).
    
    Args:
        file_path: Path to .docx file
    
    Returns:
        Extracted text with basic formatting preserved
    """
    import unicodedata
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = unicodedata.normalize('NFC', paragraph.text.strip())
            if text:
                # Preserve heading styles
                if paragraph.style and paragraph.style.name and 'heading' in paragraph.style.name.lower():
                    text_parts.append('**' + text + '**')
                else:
                    text_parts.append(text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    row_texts.append(cell_text)
                if any(row_texts):
                    text_parts.append(' | '.join(row_texts))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def get_docx_metadata(file_path: str) -> Dict:
    """
    Extract metadata from DOCX file.
    
    Returns:
        Dict with title, author, created, modified, etc.
    """
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        print(f"Metadata extraction error: {e}")
        return {}
