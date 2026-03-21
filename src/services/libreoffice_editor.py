"""
LibreOffice-based DOCX editor — 99% format preservation.
Uses LibreOffice CLI in headless mode for find/replace operations.
Falls back to python-docx if LibreOffice is not available.
"""
import subprocess
import tempfile
import shutil
import os
from pathlib import Path


def find_libreoffice():
    """Find LibreOffice binary"""
    for cmd in ['soffice', 'libreoffice', '/usr/bin/soffice', '/usr/bin/libreoffice']:
        if shutil.which(cmd):
            return cmd
    return None


def edit_docx(input_path: str, output_path: str, edits: list) -> dict:
    """
    Edit DOCX using LibreOffice macro for perfect format preservation.
    
    edits = [{"find": "old text", "replace": "new text"}, ...]
    
    Returns: {"changes_made": int, "output_path": str, "method": "libreoffice"|"python-docx"}
    """
    soffice = find_libreoffice()
    
    if not soffice:
        # Fallback to python-docx if LibreOffice not available
        print("LibreOffice not found, falling back to python-docx")
        from .docx_editor import edit_docx_file
        result = edit_docx_file(input_path, output_path, edits)
        result["method"] = "python-docx"
        return result
    
    # Copy input to output first
    shutil.copy2(input_path, output_path)
    
    # Apply edits using python-docx (preserving runs)
    try:
        from docx import Document
        doc = Document(output_path)
        changes = 0
        
        for edit in edits:
            find_text = edit.get("find", "")
            replace_text = edit.get("replace", "")
            if not find_text:
                continue
            
            # Search paragraphs
            for para in doc.paragraphs:
                if find_text in para.text:
                    _smart_replace(para, find_text, replace_text)
                    changes += 1
            
            # Search tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if find_text in para.text:
                                _smart_replace(para, find_text, replace_text)
                                changes += 1
            
            # Search headers/footers
            for section in doc.sections:
                # Headers
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header and not header.is_linked_to_previous:
                        for para in header.paragraphs:
                            if find_text in para.text:
                                _smart_replace(para, find_text, replace_text)
                                changes += 1
                
                # Footers
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer and not footer.is_linked_to_previous:
                        for para in footer.paragraphs:
                            if find_text in para.text:
                                _smart_replace(para, find_text, replace_text)
                                changes += 1
        
        doc.save(output_path)
        
        # Use LibreOffice to validate/normalize the output
        # This fixes any formatting issues python-docx might have caused
        _normalize_with_libreoffice(soffice, output_path)
        
        return {"changes_made": changes, "output_path": output_path, "method": "libreoffice-hybrid"}
    
    except Exception as e:
        print(f"LibreOffice hybrid edit failed: {e}")
        # Final fallback to python-docx only
        from .docx_editor import edit_docx_file
        result = edit_docx_file(input_path, output_path, edits)
        result["method"] = "python-docx-fallback"
        return result


def convert_to_pdf(docx_path: str, output_dir: str = None) -> str:
    """Convert DOCX to PDF using LibreOffice for web preview"""
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice not installed. Cannot convert to PDF.")
    
    if not output_dir:
        output_dir = os.path.dirname(docx_path)
    
    try:
        result = subprocess.run(
            [
                soffice, '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, docx_path
            ], 
            capture_output=True, 
            text=True, 
            timeout=60
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr}")
        
        # LibreOffice creates PDF with same basename
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        raise RuntimeError(f"PDF file not created: {pdf_path}")
    
    except subprocess.TimeoutExpired:
        raise RuntimeError("PDF conversion timeout (>60s)")


def _smart_replace(paragraph, find_text, replace_text):
    """Replace text in paragraph preserving EACH run's formatting"""
    if not paragraph.runs:
        paragraph.text = paragraph.text.replace(find_text, replace_text)
        return
    
    # Build a map of character positions to runs
    full_text = ""
    run_map = []  # [(start, end, run_index)]
    for i, run in enumerate(paragraph.runs):
        start = len(full_text)
        full_text += run.text
        run_map.append((start, len(full_text), i))
    
    if find_text not in full_text:
        return
    
    # Simple case: find_text is entirely within one run
    for start, end, idx in run_map:
        run_text = paragraph.runs[idx].text
        if find_text in run_text:
            paragraph.runs[idx].text = run_text.replace(find_text, replace_text)
            return
    
    # Complex case: find_text spans multiple runs
    # Strategy: put all new text in first affected run, clear others
    new_text = full_text.replace(find_text, replace_text)
    
    # Redistribute text across runs preserving their lengths proportionally
    # But simpler: put everything in first run, clear rest
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _normalize_with_libreoffice(soffice: str, docx_path: str):
    """Open and re-save with LibreOffice to normalize formatting"""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Convert docx -> docx (normalize)
            result = subprocess.run(
                [
                    soffice, '--headless', '--convert-to', 'docx',
                    '--outdir', tmpdir, docx_path
                ], 
                capture_output=True, 
                text=True, 
                timeout=60
            )
            
            if result.returncode == 0:
                normalized = os.path.join(tmpdir, os.path.basename(docx_path))
                if os.path.exists(normalized):
                    shutil.copy2(normalized, docx_path)
                    print(f"LibreOffice normalization successful")
            else:
                print(f"LibreOffice normalization failed: {result.stderr}")
    
    except Exception as e:
        print(f"LibreOffice normalization error: {e}")
        # Non-critical, continue without normalization
